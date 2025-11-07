# parse_schedule.py — phiên bản đọc bảng 2 cột (Ngày | Công việc)
# yêu cầu: pip install python-docx

import argparse, json, re, datetime as dt
from docx import Document

# =========================
# Regex cho các thành phần
# =========================
RE_DOW       = re.compile(r"\b(Thứ\s*[2-7]|Thứ\s*(Hai|Ba|Tư|Năm|Sáu|Bảy)|CN|Chủ nhật)\b", re.IGNORECASE)
RE_DDMM      = re.compile(r"(\d{1,2})[/-](\d{1,2})(?:[/-](\d{2,4}))?")
# Chỉ khớp giờ khi có ':' hoặc 'h/H' để không lẫn với ngày (vd 03/4)
RE_TIME      = re.compile(r"\b(\d{1,2})(?:(?::|[hH])\s?(\d{2})?)\b")
RE_LOC       = re.compile(r"\b(tại|địa điểm)\b[:\-]?\s*(?P<loc>[^:\n]+?)(?:[:.;]|$)", re.IGNORECASE)
RE_TP        = re.compile(r"^\s*(TP|Thành phần)\s*[:\-]\s*(.+)$", re.IGNORECASE)
RE_ALLDAY    = re.compile(r"\b(cả\s*ngày)\b", re.IGNORECASE)
RE_ONLINE    = re.compile(r"\b(trực tuyến|online)\b", re.IGNORECASE)
RE_PLATFORM  = re.compile(r"(MS\s*Teams?|MSTeams?|MSTeam|Zoom|Google\s*Meet|Google\s*Meeting)", re.IGNORECASE)

# =========================
# Chuẩn hoá / Utilities
# =========================
# Bản đồ chuẩn hoá tên platform (chấp nhận có/không dấu cách)
PLATFORM_MAP = {
    "msteams": "MS Teams",
    "msteam": "MS Teams",
    "ms teams": "MS Teams",
    "zoom": "Zoom",
    "googlemeet": "Google Meet",
    "google meeting": "Google Meet",
    "google meet": "Google Meet",
}

def norm(s: str) -> str:
    """Rút gọn khoảng trắng dư thừa."""
    return re.sub(r"\s+", " ", (s or "")).strip()

def _platform_key(s: str) -> str:
    """Chuẩn hoá key platform để tra map (hạ chữ + loại bỏ khoảng trắng thừa)."""
    return re.sub(r"\s+", " ", (s or "")).strip().lower().replace(" ", "")

def title_case_location(s: str) -> str:
    """
    Viết hoa mỗi từ cho location theo kiểu tự nhiên, giữ nguyên viết tắt phổ biến.
    VD: 'phòng họp số 1 nhà i' -> 'Phòng Họp Số 1 Nhà I', 'ms teams' -> 'MS Teams'
    """
    if not s:
        return s
    # Tách nhanh platform (nếu có) để không bị camel-case sai
    m = RE_PLATFORM.search(s)
    plat = None
    if m:
        plat_key = _platform_key(m.group(0))
        plat = PLATFORM_MAP.get(plat_key, m.group(0))
        s = RE_PLATFORM.sub("", s).strip()

    words = []
    for w in s.split():
        lw = w.lower()
        if lw in {"bgH".lower(), "cthđ".lower(), "hđ".lower(), "đ/c"}:
            words.append(w.upper())
        elif len(w) == 1 and w.isalpha():
            # 'i' -> 'I'
            words.append(w.upper())
        else:
            words.append(w[:1].upper() + w[1:])
    out = " ".join(words).strip()
    if plat:
        out = norm(f"{out} {plat}")
    return out

def parse_header_date(cell_text: str, default_year: int | None) -> tuple[str | None, str | None]:
    """Trích xuất thứ + ngày từ ô cột ngày."""
    t = norm(cell_text)
    m_dow = RE_DOW.search(t)
    dow = norm(m_dow.group(0)) if m_dow else None

    m = RE_DDMM.search(t)
    if not m:
        return dow, None

    d, mth, yr = int(m.group(1)), int(m.group(2)), m.group(3)
    if yr is None:
        y = default_year or dt.date.today().year
    else:
        y = int(yr)
        if y < 100:
            y = 2000 + y

    try:
        ddmmyy = dt.date(y, mth, d).strftime("%d/%m/%Y")
    except ValueError:
        ddmmyy = None
    return dow, ddmmyy

def split_events(text: str) -> list[list[str]]:
    """
    Tách nội dung thành từng block sự kiện theo bullet (*, •).
    Gộp các dòng tiếp theo (kể cả dòng 'TP: ...') cho đến khi gặp bullet mới.
    """
    lines = [l for l in (text or "").splitlines() if norm(l)]
    events, cur = [], []
    for line in lines:
        if line.lstrip().startswith(("*", "•")) and cur:
            events.append(cur)
            cur = [line]
        else:
            cur.append(line)
    if cur:
        events.append(cur)
    return events

def _clean_title(text: str) -> str:
    """Làm sạch tiêu đề: bỏ bullet, bỏ tiền tố giờ, bỏ 'TP:' kéo theo."""
    t = re.sub(r"^[*•]\s*", "", text).strip()
    t = re.sub(r"^\d{1,2}(?:(?::|[hH])\s?\d{0,2})?\s*(?:-|–|—)?\s*", "", t)  # bỏ giờ đầu câu
    t = re.sub(r"\s*(TP|Thành phần)\s*[:\-].*$", "", t, flags=re.IGNORECASE).strip()
    # bỏ cụm "qua phần mềm"
    t = re.sub(r"\bqua\s+phần\s+mềm\s*:?\s*", "", t, flags=re.IGNORECASE).strip()
    return norm(t)

def parse_event_lines(lines: list[str]) -> dict:
    """Trích xuất các trường từ một block sự kiện."""
    full = norm(" ".join(lines))

    # 1) Cả ngày?
    is_allday = bool(RE_ALLDAY.search(full))

    # 2) Thời gian
    start = end = None
    if not is_allday:
        times = RE_TIME.findall(full)  # [(h, m?), ...]
        if times:
            h1, m1 = times[0][0], times[0][1] or "00"
            start = f"{int(h1):02d}:{int(m1):02d}"
            if len(times) >= 2:
                h2, m2 = times[1][0], times[1][1] or "00"
                end = f"{int(h2):02d}:{int(m2):02d}"

    # 3) Địa điểm + nguồn tạo title
    loc = None
    title_source = full
    mloc = RE_LOC.search(full)
    if mloc:
        loc = norm(mloc.group("loc"))
        loc = re.sub(r"[,\-–—]\s*$", "", loc)
        after_loc = full[mloc.end():]
        title_source = norm(after_loc)
    else:
        # Không có "tại|địa điểm", thử nhận diện 'trực tuyến'
        if RE_ONLINE.search(full):
            plat_match = RE_PLATFORM.search(full)
            if plat_match:
                plat_std = PLATFORM_MAP.get(_platform_key(plat_match.group(0)), plat_match.group(0))
                loc = f"Trực tuyến qua {plat_std}"
            else:
                loc = "Trực tuyến"
            # bỏ cụm online khỏi title
            title_source = norm(RE_ONLINE.sub("", full))
            title_source = norm(RE_PLATFORM.sub("", title_source))

    # 4) Thành phần
    participants = None
    for l in lines:
        mtp = RE_TP.match(l)
        if mtp:
            participants = norm(mtp.group(2))
            break

    # 5) Tiêu đề
    title = _clean_title(title_source)
    if not title:
        title = _clean_title(full)

    # 6) Chuẩn hoá location
    if loc:
        loc = title_case_location(loc)

    return {
        "start": start,
        "end": end,
        "location": loc,
        "participants": participants,
        "title": title,
        "raw": full,
    }

def infer_year_from_doc(doc: Document) -> int | None:
    """Dò năm mặc định từ bất kỳ ngày có năm trong file."""
    for p in doc.paragraphs:
        m = RE_DDMM.search(p.text)
        if m and m.group(3):
            y = int(m.group(3))
            return (2000 + y) if y < 100 else y
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                m = RE_DDMM.search(cell.text)
                if m and m.group(3):
                    y = int(m.group(3))
                    return (2000 + y) if y < 100 else y
    return None

def parse_docx_as_table(path: str, default_year: int | None) -> list[dict]:
    """Parse toàn bộ docx dạng bảng 2 cột (Ngày | Công việc)."""
    doc = Document(path)
    results = []
    for t in doc.tables:
        for r in t.rows:
            if len(r.cells) < 2:
                continue
            day_text, work_text = r.cells[0].text, r.cells[1].text
            dow, the_date = parse_header_date(day_text, default_year)
            # Bỏ các hàng tiêu đề/rỗng không có DOW/DATE
            if not (dow or the_date):
                continue

            blocks = split_events(work_text)
            if not blocks:
                blocks = [[work_text]]

            for b in blocks:
                ev = parse_event_lines(b)
                ev["date"] = the_date
                ev["dow"]  = norm(dow) if dow else None
                results.append(ev)
    return results

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True, help=".docx path")
    ap.add_argument("--out", required=True, help="output jsonl path")
    ap.add_argument("--year", type=int, default=None, help="năm mặc định nếu cột ngày không ghi năm (vd 2025)")
    args = ap.parse_args()

    doc = Document(args.input)
    default_year = args.year or infer_year_from_doc(doc) or dt.date.today().year
    events = parse_docx_as_table(args.input, default_year)

    with open(args.out, "w", encoding="utf-8") as f:
        for ev in events:
            f.write(json.dumps(ev, ensure_ascii=False) + "\n")

    print(f"[OK] Parsed {len(events)} events → {args.out} (default_year={default_year})")